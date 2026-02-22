from __future__ import annotations

import io
import datetime
from typing import Dict, Any, Tuple

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _hex_to_rgb(hexstr: str) -> Tuple[int, int, int]:
    hexstr = hexstr.lstrip("#")
    return tuple(int(hexstr[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex.lstrip("#"))
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size_pt: int = 9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*_hex_to_rgb(color))


def _add_colored_heading(doc: Document, text: str, level: int, color_hex: str):
    p = doc.add_paragraph()
    style_name = f"CustomHeading{level}"
    existing = [s.name for s in doc.styles]
    if style_name not in existing:
        st = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles[f"Heading {min(level, 9)}"]
        st.font.color.rgb = RGBColor(*_hex_to_rgb(color_hex))
        st.font.bold = True
    p.style = doc.styles[style_name]
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(*_hex_to_rgb(color_hex))
    run.bold = True
    return p


def _set_run_font(run, name: str = "Montserrat", size_pt: int | None = None):
    """Apply Montserrat (or fallback Calibri) to a run."""
    run.font.name = name
    if size_pt:
        run.font.size = Pt(size_pt)
    # Word theme-font override
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _el
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(_qn("w:rFonts"))
    if rFonts is None:
        rFonts = _el("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(_qn("w:ascii"),    name)
    rFonts.set(_qn("w:hAnsi"),    name)
    rFonts.set(_qn("w:cs"),       name)


def build_docx(report: Dict[str, Any], primary_hex: str, accent_hex: str) -> bytes:
    doc = Document()

    # Set default font to Montserrat for the whole document
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _el
    style = doc.styles["Normal"]
    style.font.name = "Montserrat"
    rPr = style.element.get_or_add_rPr()
    rFonts = _el("w:rFonts")
    rFonts.set(_qn("w:ascii"), "Montserrat")
    rFonts.set(_qn("w:hAnsi"), "Montserrat")
    rPr.insert(0, rFonts)

    # Cover block
    title = doc.add_paragraph()
    r = title.add_run("INPECO  ·  Tender Intake Report")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(*_hex_to_rgb(primary_hex))
    _set_run_font(r, size_pt=22)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    r2 = subtitle.add_run(f"{report.get('tender_title', '')} — {report.get('tender_date', '')}")
    r2.italic = True
    r2.font.size = Pt(11)
    _set_run_font(r2, size_pt=11)

    gen_p = doc.add_paragraph(f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _set_run_font(gen_p.runs[0] if gen_p.runs else gen_p.add_run(""))

    # 1 Executive summary
    _add_colored_heading(doc, "1. Executive Summary", 1, primary_hex)
    for line in report.get("executive_summary", []):
        doc.add_paragraph(line, style="List Bullet")

    # 2 Deadlines
    _add_colored_heading(doc, "2. Key Deadlines & Milestones", 1, primary_hex)
    deadlines = report.get("deadlines", [])
    if deadlines:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        for i, h in enumerate(["Milestone", "When", "Evidence"]):
            _set_cell_text(hdr[i], h, bold=True, color="ffffff")
            _set_cell_shading(hdr[i], primary_hex)

        for d in deadlines:
            if not isinstance(d, dict):
                continue
            row = table.add_row().cells
            _set_cell_text(row[0], str(d.get("milestone", "")))
            _set_cell_text(row[1], str(d.get("when", "")))
            _set_cell_text(row[2], str(d.get("evidence", "")))
    else:
        doc.add_paragraph("No explicit deadlines detected.", style="List Bullet")

    # 3 Requirements
    _add_colored_heading(doc, "3. Requirements & Constraints (Extracted)", 1, primary_hex)
    reqs = report.get("requirements", {})
    if reqs:
        for cat, items in reqs.items():
            _add_colored_heading(doc, cat.replace("_", " ").title(), 2, accent_hex)
            for it in items:
                p = doc.add_paragraph(style="List Bullet")
                if isinstance(it, dict):
                    p.add_run(str(it.get("text", "")))
                    ev = it.get("evidence")
                    if ev:
                        p.add_run(f" ({ev})").italic = True
                else:
                    p.add_run(str(it))
    else:
        doc.add_paragraph("No structured requirements extracted.", style="List Bullet")

    # 4 Deliverables
    _add_colored_heading(doc, "4. Deliverables to Prepare (Pre-Bid)", 1, primary_hex)
    for item in report.get("deliverables", []):
        doc.add_paragraph(str(item), style="List Bullet")

    # 4b Showstoppers
    showstoppers = report.get("showstoppers", [])
    if showstoppers:
        _add_colored_heading(doc, "4b. Showstoppers — Immediate NO-GO Flags", 1, accent_hex)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        for i, h in enumerate(["ID", "Description", "Document Reference", "Impact"]):
            _set_cell_text(hdr[i], h, bold=True, color="ffffff")
            _set_cell_shading(hdr[i], accent_hex)
        for ss in showstoppers:
            if not isinstance(ss, dict):
                continue
            row = table.add_row().cells
            _set_cell_text(row[0], str(ss.get("id", "")))
            _set_cell_text(row[1], str(ss.get("description", "")))
            _set_cell_text(row[2], str(ss.get("document_ref", ss.get("evidence", ""))))
            _set_cell_text(row[3], str(ss.get("impact", "")))

    # 5 Risks
    _add_colored_heading(doc, "5. Risk Register", 1, primary_hex)
    risks = report.get("risks", [])
    if risks:
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        headers = ["ID", "Risk", "Category", "Level", "Score (0-100)", "Document Reference"]
        for i, h in enumerate(headers):
            _set_cell_text(hdr[i], h, bold=True, color="ffffff")
            _set_cell_shading(hdr[i], primary_hex)

        level_colors = {"High": "e74c3c", "Medium": "F7941D", "Low": "2ecc71"}
        risks = [r for r in risks if isinstance(r, dict)]
        for rsk in sorted(risks, key=lambda r: r.get("score", 0), reverse=True):
            row = table.add_row().cells
            _set_cell_text(row[0], str(rsk.get("id", "")))
            _set_cell_text(row[1], str(rsk.get("risk", "")))
            _set_cell_text(row[2], str(rsk.get("category", "")))
            lvl = str(rsk.get("level", ""))
            _set_cell_text(row[3], lvl, bold=True, color=level_colors.get(lvl, "003865"))
            _set_cell_text(row[4], str(rsk.get("score", "")))
            doc_ref = str(rsk.get("document_ref", rsk.get("evidence", "")))
            _set_cell_text(row[5], doc_ref)
    else:
        doc.add_paragraph("No major risks identified.", style="List Bullet")

    # 6 Go/No-Go
    _add_colored_heading(doc, "6. Go / No-Go Recommendation", 1, primary_hex)
    gn = report.get("go_nogo", {})
    if not isinstance(gn, dict):
        gn = {}
    doc.add_paragraph(f"Overall complexity score: {gn.get('score','?')} / 100")
    rec_p = doc.add_paragraph()
    rec_run = rec_p.add_run(f"Recommendation: {gn.get('recommendation','')}")
    rec_run.bold = True
    rec_run.font.size = Pt(14)
    rec_run.font.color.rgb = RGBColor(*_hex_to_rgb(accent_hex if gn.get("recommendation") != "GO" else primary_hex))

    if gn.get("rationale"):
        doc.add_paragraph(str(gn["rationale"]))

    # Save to memory (Streamlit Cloud friendly)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
